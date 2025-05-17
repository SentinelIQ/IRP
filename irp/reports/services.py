import logging
import os
import uuid
import json
import shutil
from datetime import datetime
from typing import Dict, List, Optional, Any, Union
from pathlib import Path

from django.conf import settings
from django.utils import timezone
from django.template import Template, Context
from django.utils.text import slugify
import markdown
import jinja2
from docx import Document
from docx.shared import Pt, RGBColor, Inches
import weasyprint

from .models import ReportTemplate, GeneratedReport, ScheduledReport
from irp.cases.models import Case, CaseObservable, CaseMitreTechnique, Task, CaseComment, CaseAttachment
from irp.observables.models import Observable
from irp.timeline.models import TimelineEvent

logger = logging.getLogger(__name__)

class ReportService:
    """
    Serviço para geração de relatórios de casos
    """
    
    # Diretório para armazenar os relatórios gerados
    REPORTS_DIR = Path(settings.MEDIA_ROOT) / "reports"
    
    @classmethod
    def ensure_reports_dir(cls):
        """
        Garante que o diretório de relatórios existe
        """
        os.makedirs(cls.REPORTS_DIR, exist_ok=True)
    
    @classmethod
    def generate_report(
        cls,
        case: Case,
        template: Optional[ReportTemplate] = None,
        output_format: Optional[str] = None,
        sections: Optional[List[str]] = None,
        include_attachments: bool = False,
        custom_header: Optional[str] = None,
        custom_footer: Optional[str] = None,
        generated_by=None
    ) -> GeneratedReport:
        """
        Gera um relatório para um caso
        """
        # Garantir que o diretório existe
        cls.ensure_reports_dir()
        
        # Se não foi especificado um template, usar o formato diretamente
        if not template and not output_format:
            output_format = 'MARKDOWN'  # Default
        elif template and not output_format:
            output_format = template.output_format
            
        # Normalizar o formato
        output_format = output_format.upper()
        
        # Criar registro do relatório
        report = GeneratedReport.objects.create(
            case=case,
            template=template,
            generated_by=generated_by,
            output_format=output_format,
            status='GENERATING',
            included_sections=sections or []
        )
        
        try:
            # Adicionar 'attachments' às seções se include_attachments for True
            if include_attachments and sections and 'attachments' not in sections:
                sections.append('attachments')
            elif include_attachments and not sections:
                sections = ['attachments']
                
            # Coletar dados do caso
            case_data = cls._collect_case_data(case, sections)
            
            # Criar diretório para o relatório específico (para anexos)
            report_dir = cls.REPORTS_DIR / f"report_{case.case_id}_{report.report_id}"
            os.makedirs(report_dir, exist_ok=True)
            
            # Criar diretório de anexos, se necessário
            attachments_dir = None
            if include_attachments and case_data.get('attachments'):
                attachments_dir = report_dir / "attachments"
                os.makedirs(attachments_dir, exist_ok=True)
                
                # Copiar anexos para o diretório do relatório
                for attachment in case_data.get('attachments', []):
                    src_path = attachment.get('file_path')
                    if src_path and os.path.exists(src_path):
                        filename = attachment.get('filename')
                        dst_path = attachments_dir / filename
                        shutil.copy2(src_path, dst_path)
                        # Atualizar caminho para referência relativa ao relatório
                        attachment['report_relative_path'] = f"attachments/{filename}"
            
            # Gerar o relatório no formato especificado
            if output_format == 'MARKDOWN':
                file_path, file_size = cls._generate_markdown_report(
                    case_data, template, custom_header, custom_footer, report.report_id, include_attachments
                )
            elif output_format == 'DOCX':
                file_path, file_size = cls._generate_docx_report(
                    case_data, template, custom_header, custom_footer, report.report_id, include_attachments
                )
            elif output_format == 'PDF':
                file_path, file_size = cls._generate_pdf_report(
                    case_data, template, custom_header, custom_footer, report.report_id, include_attachments
                )
            else:
                raise ValueError(f"Formato de saída não suportado: {output_format}")
            
            # Atualizar registro com o caminho do arquivo e tamanho
            report.file_path = file_path
            report.file_size = file_size
            report.status = 'COMPLETED'
            report.save()
            
            return report
            
        except Exception as e:
            # Registrar erro
            report.status = 'FAILED'
            report.error_message = str(e)
            report.save()
            logger.error(f"Erro ao gerar relatório para caso {case.case_id}: {str(e)}")
            return report
    
    @classmethod
    def _collect_case_data(cls, case: Case, sections: Optional[List[str]] = None) -> Dict:
        """
        Coleta todos os dados relevantes do caso para o relatório
        """
        data = {
            'case': {
                'id': str(case.case_id),
                'title': case.title,
                'description': case.description,
                'severity': case.severity.name if case.severity else None,
                'severity_color': case.severity.color_code if case.severity else None,
                'status': case.status.name if case.status else None,
                'status_color': case.status.color_code if case.status else None,
                'created_at': case.created_at,
                'updated_at': case.updated_at,
                'closed_at': case.closed_at,
                'tags': case.tags,
                'assignee': {
                    'username': case.assignee.username if case.assignee else None,
                    'full_name': f"{case.assignee.first_name} {case.assignee.last_name}".strip() if case.assignee else None
                },
                'reporter': {
                    'username': case.reporter.username if case.reporter else None,
                    'full_name': f"{case.reporter.first_name} {case.reporter.last_name}".strip() if case.reporter else None
                },
                'organization': {
                    'id': case.organization.organization_id,
                    'name': case.organization.name
                }
            },
            'metadata': {
                'generated_at': timezone.now(),
                'sections': sections or []
            }
        }
        
        # Incluir seções específicas se solicitado
        if not sections or 'observables' in sections:
            data['observables'] = cls._collect_case_observables(case)
            
        if not sections or 'timeline' in sections:
            data['timeline'] = cls._collect_case_timeline(case)
            
        if not sections or 'mitre_techniques' in sections:
            data['mitre_techniques'] = cls._collect_case_mitre_techniques(case)
            
        if not sections or 'tasks' in sections:
            data['tasks'] = cls._collect_case_tasks(case)
            
        if not sections or 'comments' in sections:
            data['comments'] = cls._collect_case_comments(case)
            
        if not sections or 'custom_fields' in sections:
            data['custom_fields'] = cls._collect_case_custom_fields(case)
            
        if not sections or 'alerts' in sections:
            data['alerts'] = cls._collect_case_alerts(case)
            
        if not sections or 'attachments' in sections:
            data['attachments'] = cls._collect_case_attachments(case)
        
        return data
    
    @staticmethod
    def _collect_case_observables(case: Case) -> List[Dict]:
        """
        Coleta os observáveis do caso
        """
        result = []
        case_observables = CaseObservable.objects.select_related('observable', 'observable__type', 'observable__tlp_level').filter(case=case)
        
        for case_obs in case_observables:
            obs = case_obs.observable
            result.append({
                'id': str(obs.observable_id),
                'value': obs.value,
                'type': obs.type.name if obs.type else None,
                'description': obs.description,
                'is_ioc': obs.is_ioc,
                'tlp_level': obs.tlp_level.name if obs.tlp_level else None,
                'tags': obs.tags,
                'sighted_at': case_obs.sighted_at
            })
        
        return result
    
    @staticmethod
    def _collect_case_timeline(case: Case) -> List[Dict]:
        """
        Coleta os eventos da timeline do caso
        """
        result = []
        timeline_events = TimelineEvent.objects.filter(case=case).order_by('occurred_at')
        
        for event in timeline_events:
            result.append({
                'id': str(event.event_id),
                'event_type': event.event_type,
                'description': event.description,
                'occurred_at': event.occurred_at,
                'actor': event.actor.username if event.actor else None,
                'metadata': event.metadata
            })
        
        return result
    
    @staticmethod
    def _collect_case_mitre_techniques(case: Case) -> List[Dict]:
        """
        Coleta as técnicas MITRE ATT&CK do caso
        """
        result = []
        case_techniques = CaseMitreTechnique.objects.select_related('technique').filter(case=case)
        
        for case_tech in case_techniques:
            tech = case_tech.technique
            result.append({
                'id': tech.technique_id,
                'name': tech.name,
                'description': tech.description,
                'url': tech.url,
                'is_subtechnique': tech.is_subtechnique,
                'parent_technique': tech.parent_technique.technique_id if tech.parent_technique else None,
                'tactics': [{'id': t.tactic_id, 'name': t.name} for t in tech.tactics.all()],
                'context_notes': case_tech.context_notes,
                'linked_at': case_tech.linked_at,
                'linked_by': case_tech.linked_by.username if case_tech.linked_by else None
            })
        
        return result
    
    @staticmethod
    def _collect_case_tasks(case: Case) -> List[Dict]:
        """
        Coleta as tarefas do caso
        """
        result = []
        tasks = Task.objects.select_related('status', 'assignee').filter(case=case).order_by('order')
        
        for task in tasks:
            result.append({
                'id': str(task.task_id),
                'title': task.title,
                'description': task.description,
                'status': task.status.name if task.status else None,
                'assignee': task.assignee.username if task.assignee else None,
                'due_date': task.due_date,
                'completed_at': task.completed_at,
                'created_at': task.created_at,
                'updated_at': task.updated_at
            })
        
        return result
    
    @staticmethod
    def _collect_case_comments(case: Case) -> List[Dict]:
        """
        Coleta os comentários do caso
        """
        result = []
        comments = CaseComment.objects.select_related('user').filter(case=case).order_by('created_at')
        
        for comment in comments:
            result.append({
                'id': str(comment.comment_id),
                'text': comment.text,
                'user': comment.user.username if comment.user else None,
                'user_full_name': f"{comment.user.first_name} {comment.user.last_name}".strip() if comment.user else None,
                'created_at': comment.created_at
            })
        
        return result
    
    @staticmethod
    def _collect_case_custom_fields(case: Case) -> List[Dict]:
        """
        Coleta os campos personalizados do caso
        """
        result = []
        
        try:
            from irp.cases.models import CaseCustomFieldValue
            
            custom_fields = CaseCustomFieldValue.objects.select_related('field_definition').filter(case=case)
            
            for cf in custom_fields:
                result.append({
                    'id': str(cf.value_id),
                    'name': cf.field_definition.name if cf.field_definition else "Unknown Field",
                    'field_type': cf.field_definition.field_type if cf.field_definition else "text",
                    'value': cf.value
                })
                
        except ImportError:
            # Se o modelo não existir, retornar lista vazia
            pass
            
        return result
    
    @staticmethod
    def _collect_case_alerts(case: Case) -> List[Dict]:
        """
        Coleta os alertas relacionados ao caso
        """
        result = []
        
        try:
            from irp.alerts.models import Alert
            
            alerts = Alert.objects.filter(related_cases__case_id=case.case_id)
            
            for alert in alerts:
                result.append({
                    'id': str(alert.alert_id),
                    'title': alert.title,
                    'description': alert.description,
                    'severity': alert.severity.name if alert.severity else None,
                    'status': alert.status.name if alert.status else None,
                    'source_system': alert.source_system,
                    'created_at': alert.created_at
                })
                
        except ImportError:
            # Se o modelo não existir, retornar lista vazia
            pass
            
        return result
    
    @staticmethod
    def _collect_case_attachments(case: Case) -> List[Dict]:
        """
        Coleta os anexos do caso
        """
        result = []
        
        attachments = CaseAttachment.objects.filter(case=case)
        
        for attachment in attachments:
            result.append({
                'id': str(attachment.attachment_id) if hasattr(attachment, 'attachment_id') else str(uuid.uuid4()),
                'filename': attachment.filename,
                'file_path': attachment.file_path,
                'content_type': attachment.content_type,
                'file_size': attachment.file_size,
                'description': attachment.description,
                'uploaded_at': attachment.uploaded_at,
                'uploaded_by': {
                    'username': attachment.uploaded_by.username if attachment.uploaded_by else None,
                    'full_name': f"{attachment.uploaded_by.first_name} {attachment.uploaded_by.last_name}".strip() if attachment.uploaded_by else None
                }
            })
            
        return result
    
    @classmethod
    def _generate_markdown_report(
        cls,
        case_data: Dict,
        template: Optional[ReportTemplate],
        custom_header: Optional[str],
        custom_footer: Optional[str],
        report_id: uuid.UUID,
        include_attachments: bool = False
    ) -> tuple:
        """
        Gera um relatório em formato Markdown
        """
        # Determinar conteúdo do template
        if template and template.template_content:
            template_content = template.template_content
        else:
            template_content = cls._get_default_markdown_template(case_data)
            
        # Adicionar cabeçalho e rodapé personalizados
        if custom_header:
            template_content = f"{custom_header}\n\n{template_content}"
            
        if custom_footer:
            template_content = f"{template_content}\n\n{custom_footer}"
            
        # Renderizar template usando Jinja2
        env = jinja2.Environment()
        template_obj = env.from_string(template_content)
        rendered_content = template_obj.render(**case_data)
        
        # Salvar arquivo
        filename = f"report_{case_data['case']['id']}_{report_id}.md"
        file_path = os.path.join(cls.REPORTS_DIR, filename)
        
        with open(file_path, 'w', encoding='utf-8') as f:
            f.write(rendered_content)
            
        # Obter tamanho do arquivo
        file_size = os.path.getsize(file_path)
        
        return file_path, file_size
    
    @classmethod
    def _generate_docx_report(
        cls,
        case_data: Dict,
        template: Optional[ReportTemplate],
        custom_header: Optional[str],
        custom_footer: Optional[str],
        report_id: uuid.UUID,
        include_attachments: bool = False
    ) -> tuple:
        """
        Gera um relatório em formato DOCX (Microsoft Word)
        """
        # Criar documento Word
        doc = Document()
        
        # Configurar estilos
        style = doc.styles['Normal']
        style.font.name = 'Calibri'
        style.font.size = Pt(11)
        
        # Título do documento
        case_title = case_data['case']['title']
        doc.add_heading(f"Case Report: {case_title}", level=1)
        
        # Informações básicas do caso
        doc.add_heading("Case Information", level=2)
        
        case_info = doc.add_table(rows=1, cols=2)
        case_info.style = 'Table Grid'
        case_info.allow_autofit = True
        
        hdr_cells = case_info.rows[0].cells
        hdr_cells[0].text = "Field"
        hdr_cells[1].text = "Value"
        
        case = case_data['case']
        
        # Adicionar linhas com informações do caso
        for field, value in [
            ("Case ID", case['id']),
            ("Title", case['title']),
            ("Status", case['status']),
            ("Severity", case['severity']),
            ("Created At", case['created_at'].strftime('%Y-%m-%d %H:%M:%S')),
            ("Updated At", case['updated_at'].strftime('%Y-%m-%d %H:%M:%S')),
            ("Closed At", case['closed_at'].strftime('%Y-%m-%d %H:%M:%S') if case['closed_at'] else "N/A"),
            ("Assignee", case['assignee']['full_name'] if case['assignee']['full_name'] else "Unassigned"),
            ("Reporter", case['reporter']['full_name'] if case['reporter']['full_name'] else "Unknown"),
            ("Organization", case['organization']['name']),
            ("Tags", ", ".join(case['tags']) if case['tags'] else "None")
        ]:
            row_cells = case_info.add_row().cells
            row_cells[0].text = field
            row_cells[1].text = str(value)
        
        # Descrição do caso
        doc.add_heading("Description", level=2)
        doc.add_paragraph(case['description'])
        
        # Adicionar seções específicas
        sections = case_data['metadata'].get('sections', [])
        
        if not sections or 'observables' in sections:
            cls._add_observables_to_docx(doc, case_data.get('observables', []))
            
        if not sections or 'timeline' in sections:
            cls._add_timeline_to_docx(doc, case_data.get('timeline', []))
            
        if not sections or 'mitre_techniques' in sections:
            cls._add_mitre_techniques_to_docx(doc, case_data.get('mitre_techniques', []))
            
        if not sections or 'tasks' in sections:
            cls._add_tasks_to_docx(doc, case_data.get('tasks', []))
            
        if not sections or 'comments' in sections:
            cls._add_comments_to_docx(doc, case_data.get('comments', []))
            
        if (not sections or 'attachments' in sections) and include_attachments:
            cls._add_attachments_to_docx(doc, case_data.get('attachments', []))
        
        # Adicionar rodapé
        doc.add_paragraph("\n\n")
        footer = doc.add_paragraph(f"Generated at: {timezone.now().strftime('%Y-%m-%d %H:%M:%S')}")
        footer.style = 'Subtitle'
        
        # Salvar arquivo
        filename = f"report_{case_data['case']['id']}_{report_id}.docx"
        file_path = os.path.join(cls.REPORTS_DIR, filename)
        
        doc.save(file_path)
        
        # Obter tamanho do arquivo
        file_size = os.path.getsize(file_path)
        
        return file_path, file_size
    
    @staticmethod
    def _add_observables_to_docx(doc, observables):
        """Adiciona seção de observáveis ao documento DOCX"""
        if not observables:
            return
            
        doc.add_heading("Observables", level=2)
        
        table = doc.add_table(rows=1, cols=4)
        table.style = 'Table Grid'
        
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = "Type"
        hdr_cells[1].text = "Value"
        hdr_cells[2].text = "TLP"
        hdr_cells[3].text = "Description"
        
        for obs in observables:
            row_cells = table.add_row().cells
            row_cells[0].text = obs.get('type', 'Unknown')
            row_cells[1].text = obs.get('value', '')
            row_cells[2].text = obs.get('tlp_level', 'N/A')
            row_cells[3].text = obs.get('description', '')
    
    @staticmethod
    def _add_timeline_to_docx(doc, timeline):
        """Adiciona seção de timeline ao documento DOCX"""
        if not timeline:
            return
            
        doc.add_heading("Timeline", level=2)
        
        table = doc.add_table(rows=1, cols=3)
        table.style = 'Table Grid'
        
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = "Time"
        hdr_cells[1].text = "Event Type"
        hdr_cells[2].text = "Description"
        
        for event in timeline:
            row_cells = table.add_row().cells
            row_cells[0].text = event.get('occurred_at').strftime('%Y-%m-%d %H:%M:%S')
            row_cells[1].text = event.get('event_type', '')
            row_cells[2].text = event.get('description', '')
    
    @staticmethod
    def _add_mitre_techniques_to_docx(doc, techniques):
        """Adiciona seção de técnicas MITRE ao documento DOCX"""
        if not techniques:
            return
            
        doc.add_heading("MITRE ATT&CK Techniques", level=2)
        
        for tech in techniques:
            doc.add_heading(f"{tech.get('id', '')}: {tech.get('name', '')}", level=3)
            
            p = doc.add_paragraph()
            p.add_run("Description: ").bold = True
            p.add_run(tech.get('description', ''))
            
            tactics = ", ".join([t.get('name', '') for t in tech.get('tactics', [])])
            p = doc.add_paragraph()
            p.add_run("Tactics: ").bold = True
            p.add_run(tactics)
            
            if tech.get('context_notes'):
                p = doc.add_paragraph()
                p.add_run("Context Notes: ").bold = True
                p.add_run(tech.get('context_notes', ''))
    
    @staticmethod
    def _add_tasks_to_docx(doc, tasks):
        """Adiciona seção de tarefas ao documento DOCX"""
        if not tasks:
            return
            
        doc.add_heading("Tasks", level=2)
        
        table = doc.add_table(rows=1, cols=4)
        table.style = 'Table Grid'
        
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = "Title"
        hdr_cells[1].text = "Status"
        hdr_cells[2].text = "Assignee"
        hdr_cells[3].text = "Due Date"
        
        for task in tasks:
            row_cells = table.add_row().cells
            row_cells[0].text = task.get('title', '')
            row_cells[1].text = task.get('status', '')
            row_cells[2].text = task.get('assignee', 'Unassigned')
            due_date = task.get('due_date')
            row_cells[3].text = due_date.strftime('%Y-%m-%d') if due_date else 'N/A'
            
            if task.get('description'):
                row = table.add_row()
                desc_cell = row.cells[0]
                desc_cell.merge(row.cells[1])
                desc_cell.merge(row.cells[2]) 
                desc_cell.merge(row.cells[3])
                desc_cell.text = task.get('description', '')
    
    @staticmethod
    def _add_comments_to_docx(doc, comments):
        """Adiciona seção de comentários ao documento DOCX"""
        if not comments:
            return
            
        doc.add_heading("Comments", level=2)
        
        for comment in comments:
            p = doc.add_paragraph()
            p.add_run(f"{comment.get('user_full_name', comment.get('user', 'Unknown'))} ").bold = True
            p.add_run(f"({comment.get('created_at').strftime('%Y-%m-%d %H:%M:%S')}): ")
            p.add_run(comment.get('text', ''))
    
    @staticmethod
    def _add_attachments_to_docx(doc, attachments):
        """Adiciona seção de anexos ao documento DOCX"""
        if not attachments:
            return
            
        doc.add_heading("Attachments", level=2)
        
        table = doc.add_table(rows=1, cols=4)
        table.style = 'Table Grid'
        
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = "Filename"
        hdr_cells[1].text = "Size"
        hdr_cells[2].text = "Uploaded By"
        hdr_cells[3].text = "Description"
        
        for attachment in attachments:
            row_cells = table.add_row().cells
            row_cells[0].text = attachment.get('filename', '')
            
            file_size = attachment.get('file_size', 0)
            # Formatar tamanho do arquivo
            if file_size < 1024:
                size_str = f"{file_size} bytes"
            elif file_size < 1024 * 1024:
                size_str = f"{file_size/1024:.1f} KB"
            else:
                size_str = f"{file_size/(1024*1024):.1f} MB"
                
            row_cells[1].text = size_str
            row_cells[2].text = attachment.get('uploaded_by', {}).get('full_name', 'Unknown')
            row_cells[3].text = attachment.get('description', '')
    
    @classmethod
    def _generate_pdf_report(
        cls,
        case_data: Dict,
        template: Optional[ReportTemplate],
        custom_header: Optional[str],
        custom_footer: Optional[str],
        report_id: uuid.UUID,
        include_attachments: bool = False
    ) -> tuple:
        """
        Gera um relatório em formato PDF
        """
        # Primeiro gerar conteúdo em markdown
        markdown_content = cls._get_default_markdown_template(case_data)
        
        if template and template.template_content:
            markdown_content = template.template_content
            
        # Adicionar cabeçalho e rodapé personalizados
        if custom_header:
            markdown_content = f"{custom_header}\n\n{markdown_content}"
            
        if custom_footer:
            markdown_content = f"{markdown_content}\n\n{custom_footer}"
            
        # Renderizar template usando Jinja2
        env = jinja2.Environment()
        template_obj = env.from_string(markdown_content)
        rendered_markdown = template_obj.render(**case_data)
        
        # Converter markdown para HTML
        html_content = markdown.markdown(rendered_markdown, extensions=['tables'])
        
        # Adicionar estilos ao HTML
        styled_html = f"""
        <!DOCTYPE html>
        <html>
        <head>
            <meta charset="utf-8">
            <title>Case Report: {case_data['case']['title']}</title>
            <style>
                body {{ font-family: Arial, sans-serif; font-size: 12px; line-height: 1.6; }}
                h1 {{ color: #333366; }}
                h2 {{ color: #336699; border-bottom: 1px solid #cccccc; }}
                h3 {{ color: #3399cc; }}
                table {{ border-collapse: collapse; width: 100%; margin-bottom: 20px; }}
                th, td {{ border: 1px solid #dddddd; padding: 8px; text-align: left; }}
                th {{ background-color: #f2f2f2; }}
                .header {{ text-align: center; margin-bottom: 30px; }}
                .footer {{ text-align: center; margin-top: 30px; font-size: 10px; color: #666666; }}
            </style>
        </head>
        <body>
            <div class="header">
                <h1>Case Report: {case_data['case']['title']}</h1>
                <p>Generated on {timezone.now().strftime('%Y-%m-%d %H:%M:%S')}</p>
            </div>
            
            {html_content}
            
            <div class="footer">
                <p>This report was generated automatically. Case ID: {case_data['case']['id']}</p>
            </div>
        </body>
        </html>
        """
        
        # Gerar o PDF usando WeasyPrint
        filename = f"report_{case_data['case']['id']}_{report_id}.pdf"
        file_path = os.path.join(cls.REPORTS_DIR, filename)
        
        pdf = weasyprint.HTML(string=styled_html).write_pdf()
        
        with open(file_path, 'wb') as f:
            f.write(pdf)
            
        # Obter tamanho do arquivo
        file_size = os.path.getsize(file_path)
        
        return file_path, file_size
    
    @staticmethod
    def _get_default_markdown_template(case_data: Dict) -> str:
        """
        Retorna um template markdown padrão
        """
        case = case_data['case']
        
        template = f"""# Case Report: {case['title']}

## Case Information

- **Case ID**: {case['id']}
- **Title**: {case['title']}
- **Status**: {case['status']}
- **Severity**: {case['severity']}
- **Created**: {case['created_at'].strftime('%Y-%m-%d %H:%M:%S')}
- **Updated**: {case['updated_at'].strftime('%Y-%m-%d %H:%M:%S')}
- **Closed**: {case['closed_at'].strftime('%Y-%m-%d %H:%M:%S') if case['closed_at'] else "N/A"}
- **Assignee**: {case['assignee']['full_name'] if case['assignee']['full_name'] else "Unassigned"}
- **Reporter**: {case['reporter']['full_name'] if case['reporter']['full_name'] else "Unknown"}
- **Organization**: {case['organization']['name']}
- **Tags**: {', '.join(case['tags']) if case['tags'] else "None"}

## Description

{case['description']}

"""

        # Adicionar seções específicas se existirem
        if 'observables' in case_data and case_data['observables']:
            template += """
## Observables

| Type | Value | TLP | Description |
|------|-------|-----|-------------|
"""
            for obs in case_data['observables']:
                template += f"| {obs.get('type', 'Unknown')} | {obs.get('value', '')} | {obs.get('tlp_level', 'N/A')} | {obs.get('description', '')} |\n"

        if 'timeline' in case_data and case_data['timeline']:
            template += """
## Timeline

| Time | Event Type | Description |
|------|------------|-------------|
"""
            for event in case_data['timeline']:
                template += f"| {event.get('occurred_at').strftime('%Y-%m-%d %H:%M:%S')} | {event.get('event_type', '')} | {event.get('description', '')} |\n"

        if 'mitre_techniques' in case_data and case_data['mitre_techniques']:
            template += """
## MITRE ATT&CK Techniques

"""
            for tech in case_data['mitre_techniques']:
                tactics = ", ".join([t.get('name', '') for t in tech.get('tactics', [])])
                template += f"### {tech.get('id', '')}: {tech.get('name', '')}\n\n"
                template += f"**Description**: {tech.get('description', '')}\n\n"
                template += f"**Tactics**: {tactics}\n\n"
                if tech.get('context_notes'):
                    template += f"**Context Notes**: {tech.get('context_notes', '')}\n\n"

        if 'tasks' in case_data and case_data['tasks']:
            template += """
## Tasks

| Title | Status | Assignee | Due Date | Description |
|-------|--------|----------|----------|-------------|
"""
            for task in case_data['tasks']:
                due_date = task.get('due_date')
                due_date_str = due_date.strftime('%Y-%m-%d') if due_date else 'N/A'
                template += f"| {task.get('title', '')} | {task.get('status', '')} | {task.get('assignee', 'Unassigned')} | {due_date_str} | {task.get('description', '')} |\n"

        if 'comments' in case_data and case_data['comments']:
            template += """
## Comments

"""
            for comment in case_data['comments']:
                template += f"**{comment.get('user_full_name', comment.get('user', 'Unknown'))}** ({comment.get('created_at').strftime('%Y-%m-%d %H:%M:%S')}): {comment.get('text', '')}\n\n"

        if 'attachments' in case_data and case_data['attachments']:
            template += """
## Attachments

| Filename | Size | Uploaded By | Description |
|----------|------|-------------|-------------|
"""
            for attachment in case_data['attachments']:
                file_size = attachment.get('file_size', 0)
                # Formatar tamanho do arquivo
                if file_size < 1024:
                    size_str = f"{file_size} bytes"
                elif file_size < 1024 * 1024:
                    size_str = f"{file_size/1024:.1f} KB"
                else:
                    size_str = f"{file_size/(1024*1024):.1f} MB"
                    
                template += f"| {attachment.get('filename', '')} | {size_str} | {attachment.get('uploaded_by', {}).get('full_name', 'Unknown')} | {attachment.get('description', '')} |\n"

        template += f"\n\n---\nGenerated at: {timezone.now().strftime('%Y-%m-%d %H:%M:%S')}"
        
        return template 

    @classmethod
    def process_scheduled_reports(cls):
        """
        Processa todos os relatórios agendados que estão pendentes
        Esta função é chamada periodicamente pelo Celery
        """
        from .models import ScheduledReport
        from irp.cases.models import Case
        from django.utils import timezone
        from django.db.models import Q
        
        now = timezone.now()
        logger.info(f"Processando relatórios agendados: {now}")
        
        # Obter todos os relatórios agendados pendentes
        pending_schedules = ScheduledReport.objects.filter(
            is_active=True,
            next_run__lte=now
        )
        
        logger.info(f"Encontrados {pending_schedules.count()} relatórios agendados pendentes")
        
        for schedule in pending_schedules:
            try:
                logger.info(f"Processando agendamento: {schedule.name} (ID: {schedule.schedule_id})")
                
                # Construir filtro de casos
                case_query = Q(organization=schedule.organization)
                
                # Adicionar filtros adicionais
                filters = schedule.case_filter
                if filters:
                    if filters.get('status'):
                        case_query &= Q(status__name__in=filters.get('status'))
                    if filters.get('severity'):
                        case_query &= Q(severity__name__in=filters.get('severity'))
                    if filters.get('tags'):
                        # Para cada tag no filtro, exigir que esteja presente no caso
                        for tag in filters.get('tags'):
                            case_query &= Q(tags__contains=[tag])
                
                # Filtrar casos de acordo com os critérios
                cases = Case.objects.filter(case_query)
                
                logger.info(f"Encontrados {cases.count()} casos correspondentes aos critérios")
                
                # Registrar última execução
                schedule.last_run = now
                
                # Calcular próxima execução
                schedule.calculate_next_run()
                schedule.save()
                
                # Gerar relatórios para cada caso
                for case in cases:
                    try:
                        # Gerar o relatório
                        report = cls.generate_report(
                            case=case,
                            template=schedule.template,
                            output_format=schedule.output_format,
                            sections=schedule.include_sections,
                            include_attachments=schedule.include_attachments,
                            custom_header=schedule.custom_header,
                            custom_footer=schedule.custom_footer,
                            generated_by=schedule.created_by
                        )
                        
                        # Enviar notificações para os usuários configurados
                        if schedule.send_email and report.status == 'COMPLETED':
                            cls._send_report_notification(schedule, report, case)
                            
                    except Exception as e:
                        logger.error(f"Erro ao gerar relatório para caso {case.case_id}: {str(e)}")
                
            except Exception as e:
                logger.error(f"Erro ao processar agendamento {schedule.schedule_id}: {str(e)}")
    
    @staticmethod
    def _send_report_notification(schedule, report, case):
        """
        Envia notificações para os usuários configurados sobre o relatório gerado
        """
        from django.core.mail import send_mail
        from django.conf import settings
        
        # Obter usuários que devem ser notificados
        users = schedule.notify_users.all()
        
        if not users:
            return
            
        # Preparar mensagem de email
        subject = f"[IRP] Relatório agendado gerado: {case.title}"
        message = f"""Olá,

Um relatório foi gerado automaticamente para o caso "{case.title}".

Detalhes do relatório:
- ID do relatório: {report.report_id}
- Formato: {report.output_format}
- Gerado em: {report.created_at.strftime('%Y-%m-%d %H:%M:%S')}

Você pode acessar e baixar este relatório através do sistema IRP.

--
Este é um email automático, por favor não responda.
"""
        
        # Enviar email para cada usuário
        for user in users:
            if user.email:
                try:
                    send_mail(
                        subject,
                        message,
                        settings.DEFAULT_FROM_EMAIL,
                        [user.email],
                        fail_silently=False,
                    )
                except Exception as e:
                    logger.error(f"Erro ao enviar email para {user.email}: {str(e)}") 