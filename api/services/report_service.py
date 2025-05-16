import logging
import os
import uuid
import json
from datetime import datetime
from typing import Dict, List, Optional, Any, Union
from pathlib import Path

from django.conf import settings
from django.utils import timezone
from django.template import Template, Context
from django.utils.text import slugify
import markdown
import jinja2
from docx import Document
from docx.shared import Pt, RGBColor, Inches
import weasyprint

from api.models import (
    Case, ReportTemplate, GeneratedReport, Observable, TimelineEvent, 
    CaseObservable, CaseMitreTechnique, Task, CaseComment
)

logger = logging.getLogger(__name__)

class ReportService:
    """
    Serviço para geração de relatórios de casos
    """
    
    # Diretório para armazenar os relatórios gerados
    REPORTS_DIR = Path(settings.MEDIA_ROOT) / "reports"
    
    @classmethod
    def ensure_reports_dir(cls):
        """
        Garante que o diretório de relatórios existe
        """
        os.makedirs(cls.REPORTS_DIR, exist_ok=True)
    
    @classmethod
    def generate_report(
        cls,
        case: Case,
        template: Optional[ReportTemplate] = None,
        output_format: Optional[str] = None,
        sections: Optional[List[str]] = None,
        include_attachments: bool = False,
        custom_header: Optional[str] = None,
        custom_footer: Optional[str] = None,
        generated_by=None
    ) -> GeneratedReport:
        """
        Gera um relatório para um caso
        """
        # Garantir que o diretório existe
        cls.ensure_reports_dir()
        
        # Se não foi especificado um template, usar o formato diretamente
        if not template and not output_format:
            output_format = 'MARKDOWN'  # Default
        elif template and not output_format:
            output_format = template.output_format
            
        # Normalizar o formato
        output_format = output_format.upper()
        
        # Criar registro do relatório
        report = GeneratedReport.objects.create(
            case=case,
            template=template,
            generated_by=generated_by,
            output_format=output_format,
            status='GENERATING',
            included_sections=sections or []
        )
        
        try:
            # Coletar dados do caso
            case_data = cls._collect_case_data(case, sections)
            
            # Gerar o relatório no formato especificado
            if output_format == 'MARKDOWN':
                file_path, file_size = cls._generate_markdown_report(
                    case_data, template, custom_header, custom_footer, report.report_id
                )
            elif output_format == 'DOCX':
                file_path, file_size = cls._generate_docx_report(
                    case_data, template, custom_header, custom_footer, report.report_id
                )
            elif output_format == 'PDF':
                file_path, file_size = cls._generate_pdf_report(
                    case_data, template, custom_header, custom_footer, report.report_id
                )
            else:
                raise ValueError(f"Formato de saída não suportado: {output_format}")
            
            # Atualizar registro com o caminho do arquivo e tamanho
            report.file_path = file_path
            report.file_size = file_size
            report.status = 'COMPLETED'
            report.save()
            
            return report
            
        except Exception as e:
            # Registrar erro
            report.status = 'FAILED'
            report.error_message = str(e)
            report.save()
            logger.error(f"Erro ao gerar relatório para caso {case.case_id}: {str(e)}")
            return report
    
    @classmethod
    def _collect_case_data(cls, case: Case, sections: Optional[List[str]] = None) -> Dict:
        """
        Coleta todos os dados relevantes do caso para o relatório
        """
        data = {
            'case': {
                'id': str(case.case_id),
                'title': case.title,
                'description': case.description,
                'severity': case.severity.name if case.severity else None,
                'severity_color': case.severity.color_code if case.severity else None,
                'status': case.status.name if case.status else None,
                'status_color': case.status.color_code if case.status else None,
                'created_at': case.created_at,
                'updated_at': case.updated_at,
                'closed_at': case.closed_at,
                'tags': case.tags,
                'assignee': {
                    'username': case.assignee.username if case.assignee else None,
                    'full_name': f"{case.assignee.first_name} {case.assignee.last_name}".strip() if case.assignee else None
                },
                'reporter': {
                    'username': case.reporter.username if case.reporter else None,
                    'full_name': f"{case.reporter.first_name} {case.reporter.last_name}".strip() if case.reporter else None
                },
                'organization': {
                    'id': case.organization.organization_id,
                    'name': case.organization.name
                }
            },
            'metadata': {
                'generated_at': timezone.now(),
                'sections': sections or []
            }
        }
        
        # Incluir seções específicas se solicitado
        if not sections or 'observables' in sections:
            data['observables'] = cls._collect_case_observables(case)
            
        if not sections or 'timeline' in sections:
            data['timeline'] = cls._collect_case_timeline(case)
            
        if not sections or 'mitre_techniques' in sections:
            data['mitre_techniques'] = cls._collect_case_mitre_techniques(case)
            
        if not sections or 'tasks' in sections:
            data['tasks'] = cls._collect_case_tasks(case)
            
        if not sections or 'comments' in sections:
            data['comments'] = cls._collect_case_comments(case)
            
        if not sections or 'custom_fields' in sections:
            data['custom_fields'] = cls._collect_case_custom_fields(case)
            
        if not sections or 'alerts' in sections:
            data['alerts'] = cls._collect_case_alerts(case)
        
        return data
    
    @staticmethod
    def _collect_case_observables(case: Case) -> List[Dict]:
        """
        Coleta os observáveis do caso
        """
        result = []
        case_observables = CaseObservable.objects.select_related('observable', 'observable__type', 'observable__tlp_level').filter(case=case)
        
        for case_obs in case_observables:
            obs = case_obs.observable
            result.append({
                'id': str(obs.observable_id),
                'value': obs.value,
                'type': obs.type.name if obs.type else None,
                'description': obs.description,
                'is_ioc': obs.is_ioc,
                'tlp_level': obs.tlp_level.name if obs.tlp_level else None,
                'tags': obs.tags,
                'sighted_at': case_obs.sighted_at
            })
        
        return result
    
    @staticmethod
    def _collect_case_timeline(case: Case) -> List[Dict]:
        """
        Coleta os eventos da timeline do caso
        """
        result = []
        timeline_events = TimelineEvent.objects.filter(case=case).order_by('occurred_at')
        
        for event in timeline_events:
            result.append({
                'id': str(event.event_id),
                'event_type': event.event_type,
                'description': event.description,
                'occurred_at': event.occurred_at,
                'actor': event.actor.username if event.actor else None,
                'metadata': event.metadata
            })
        
        return result
    
    @staticmethod
    def _collect_case_mitre_techniques(case: Case) -> List[Dict]:
        """
        Coleta as técnicas MITRE ATT&CK do caso
        """
        result = []
        case_techniques = CaseMitreTechnique.objects.select_related('technique').filter(case=case)
        
        for case_tech in case_techniques:
            tech = case_tech.technique
            result.append({
                'id': tech.technique_id,
                'name': tech.name,
                'description': tech.description,
                'url': tech.url,
                'is_subtechnique': tech.is_subtechnique,
                'parent_technique': tech.parent_technique.technique_id if tech.parent_technique else None,
                'tactics': [{'id': t.tactic_id, 'name': t.name} for t in tech.tactics.all()],
                'context_notes': case_tech.context_notes,
                'linked_at': case_tech.linked_at,
                'linked_by': case_tech.linked_by.username if case_tech.linked_by else None
            })
        
        return result
    
    @staticmethod
    def _collect_case_tasks(case: Case) -> List[Dict]:
        """
        Coleta as tarefas do caso
        """
        result = []
        tasks = Task.objects.select_related('status', 'assignee').filter(case=case).order_by('order')
        
        for task in tasks:
            result.append({
                'id': str(task.task_id),
                'title': task.title,
                'description': task.description,
                'status': task.status.name if task.status else None,
                'status_color': task.status.color_code if task.status else None,
                'assignee': task.assignee.username if task.assignee else None,
                'due_date': task.due_date,
                'created_at': task.created_at,
                'updated_at': task.updated_at
            })
        
        return result
    
    @staticmethod
    def _collect_case_comments(case: Case) -> List[Dict]:
        """
        Coleta os comentários do caso
        """
        result = []
        comments = CaseComment.objects.select_related('user').filter(case=case).order_by('created_at')
        
        for comment in comments:
            result.append({
                'id': str(comment.comment_id),
                'text': comment.comment_text,
                'user': comment.user.username if comment.user else None,
                'user_full_name': f"{comment.user.first_name} {comment.user.last_name}".strip() if comment.user else None,
                'created_at': comment.created_at
            })
        
        return result
    
    @staticmethod
    def _collect_case_custom_fields(case: Case) -> List[Dict]:
        """
        Coleta os campos customizados do caso
        """
        result = []
        custom_fields = case.custom_field_values.select_related('field_definition').all()
        
        for cf in custom_fields:
            # Determinar o valor baseado no tipo do campo
            if cf.value_text is not None:
                value = cf.value_text
            elif cf.value_number is not None:
                value = cf.value_number
            elif cf.value_boolean is not None:
                value = cf.value_boolean
            elif cf.value_date is not None:
                value = cf.value_date
            else:
                value = None
                
            result.append({
                'name': cf.field_definition.name,
                'technical_name': cf.field_definition.technical_name,
                'field_type': cf.field_definition.field_type,
                'value': value
            })
        
        return result
    
    @staticmethod
    def _collect_case_alerts(case: Case) -> List[Dict]:
        """
        Coleta os alertas associados ao caso
        """
        result = []
        alerts = case.alerts.select_related('severity', 'status').all()
        
        for alert in alerts:
            result.append({
                'id': str(alert.alert_id),
                'title': alert.title,
                'description': alert.description,
                'severity': alert.severity.name if alert.severity else None,
                'severity_color': alert.severity.color_code if alert.severity else None,
                'status': alert.status.name if alert.status else None,
                'status_color': alert.status.color_code if alert.status else None,
                'source_system': alert.source_system,
                'created_at': alert.created_at,
                'first_seen_at': alert.first_seen_at,
                'last_seen_at': alert.last_seen_at
            })
        
        return result
    
    @classmethod
    def _generate_markdown_report(
        cls,
        case_data: Dict,
        template: Optional[ReportTemplate],
        custom_header: Optional[str],
        custom_footer: Optional[str],
        report_id: uuid.UUID
    ) -> tuple:
        """
        Gera um relatório em formato Markdown
        """
        case = case_data['case']
        
        # Determinar o conteúdo do template
        if template and template.template_content:
            # Usar o template fornecido
            template_content = template.template_content
        else:
            # Usar um template padrão
            template_content = cls._get_default_markdown_template(case_data)
            
        # Adicionar cabeçalho e rodapé personalizados, se fornecidos
        if custom_header:
            template_content = f"{custom_header}\n\n{template_content}"
        if custom_footer:
            template_content = f"{template_content}\n\n{custom_footer}"
            
        # Renderizar o template com Jinja2
        j2_env = jinja2.Environment()
        j2_template = j2_env.from_string(template_content)
        rendered_content = j2_template.render(**case_data)
        
        # Salvar o arquivo
        file_name = f"{slugify(case['title'])}-{report_id}.md"
        file_path = cls.REPORTS_DIR / file_name
        
        with open(file_path, 'w', encoding='utf-8') as f:
            f.write(rendered_content)
            
        # Retornar o caminho relativo e o tamanho do arquivo
        file_size = os.path.getsize(file_path)
        relative_path = os.path.join('reports', file_name)
        
        return relative_path, file_size
    
    @classmethod
    def _generate_docx_report(
        cls,
        case_data: Dict,
        template: Optional[ReportTemplate],
        custom_header: Optional[str],
        custom_footer: Optional[str],
        report_id: uuid.UUID
    ) -> tuple:
        """
        Gera um relatório em formato DOCX
        """
        case = case_data['case']
        
        # Criar um novo documento
        doc = Document()
        
        # Adicionar cabeçalho personalizado, se fornecido
        if custom_header:
            doc.add_heading(custom_header, level=1)
            doc.add_paragraph()
            
        # Adicionar título e informações básicas
        doc.add_heading(f"Relatório de Caso: {case['title']}", level=1)
        doc.add_paragraph(f"ID do Caso: {case['id']}")
        doc.add_paragraph(f"Severidade: {case['severity']}")
        doc.add_paragraph(f"Status: {case['status']}")
        doc.add_paragraph(f"Data de Criação: {case['created_at'].strftime('%d/%m/%Y %H:%M')}")
        if case['closed_at']:
            doc.add_paragraph(f"Data de Encerramento: {case['closed_at'].strftime('%d/%m/%Y %H:%M')}")
        
        # Adicionar descrição
        doc.add_heading("Descrição", level=2)
        doc.add_paragraph(case['description'] or "Sem descrição disponível.")
        
        # Adicionar seções com base nos dados disponíveis
        sections = case_data['metadata']['sections']
        
        # Observáveis
        if 'observables' in case_data and (not sections or 'observables' in sections):
            doc.add_heading("Observáveis", level=2)
            if case_data['observables']:
                table = doc.add_table(rows=1, cols=4)
                table.style = 'Table Grid'
                hdr_cells = table.rows[0].cells
                hdr_cells[0].text = 'Valor'
                hdr_cells[1].text = 'Tipo'
                hdr_cells[2].text = 'IoC'
                hdr_cells[3].text = 'Descrição'
                
                for obs in case_data['observables']:
                    row_cells = table.add_row().cells
                    row_cells[0].text = obs['value']
                    row_cells[1].text = obs['type'] or 'N/A'
                    row_cells[2].text = 'Sim' if obs['is_ioc'] else 'Não'
                    row_cells[3].text = obs['description'] or ''
            else:
                doc.add_paragraph("Nenhum observável associado a este caso.")
        
        # Timeline
        if 'timeline' in case_data and (not sections or 'timeline' in sections):
            doc.add_heading("Timeline", level=2)
            if case_data['timeline']:
                table = doc.add_table(rows=1, cols=3)
                table.style = 'Table Grid'
                hdr_cells = table.rows[0].cells
                hdr_cells[0].text = 'Data/Hora'
                hdr_cells[1].text = 'Tipo'
                hdr_cells[2].text = 'Descrição'
                
                for event in case_data['timeline']:
                    row_cells = table.add_row().cells
                    row_cells[0].text = event['occurred_at'].strftime('%d/%m/%Y %H:%M')
                    row_cells[1].text = event['event_type']
                    row_cells[2].text = event['description']
            else:
                doc.add_paragraph("Nenhum evento de timeline registrado.")
        
        # Técnicas MITRE ATT&CK
        if 'mitre_techniques' in case_data and (not sections or 'mitre_techniques' in sections):
            doc.add_heading("Técnicas MITRE ATT&CK", level=2)
            if case_data['mitre_techniques']:
                table = doc.add_table(rows=1, cols=3)
                table.style = 'Table Grid'
                hdr_cells = table.rows[0].cells
                hdr_cells[0].text = 'ID'
                hdr_cells[1].text = 'Nome'
                hdr_cells[2].text = 'Táticas'
                
                for tech in case_data['mitre_techniques']:
                    row_cells = table.add_row().cells
                    row_cells[0].text = tech['id']
                    row_cells[1].text = tech['name']
                    row_cells[2].text = ', '.join(t['name'] for t in tech['tactics'])
            else:
                doc.add_paragraph("Nenhuma técnica MITRE ATT&CK associada.")
        
        # Tarefas
        if 'tasks' in case_data and (not sections or 'tasks' in sections):
            doc.add_heading("Tarefas", level=2)
            if case_data['tasks']:
                table = doc.add_table(rows=1, cols=4)
                table.style = 'Table Grid'
                hdr_cells = table.rows[0].cells
                hdr_cells[0].text = 'Título'
                hdr_cells[1].text = 'Status'
                hdr_cells[2].text = 'Responsável'
                hdr_cells[3].text = 'Data de Vencimento'
                
                for task in case_data['tasks']:
                    row_cells = table.add_row().cells
                    row_cells[0].text = task['title']
                    row_cells[1].text = task['status'] or 'N/A'
                    row_cells[2].text = task['assignee'] or 'Não atribuído'
                    row_cells[3].text = task['due_date'].strftime('%d/%m/%Y') if task['due_date'] else 'N/A'
            else:
                doc.add_paragraph("Nenhuma tarefa associada a este caso.")
        
        # Comentários
        if 'comments' in case_data and (not sections or 'comments' in sections):
            doc.add_heading("Comentários", level=2)
            if case_data['comments']:
                for comment in case_data['comments']:
                    p = doc.add_paragraph()
                    p.add_run(f"{comment['user_full_name'] or comment['user']} - {comment['created_at'].strftime('%d/%m/%Y %H:%M')}:").bold = True
                    doc.add_paragraph(comment['text'])
            else:
                doc.add_paragraph("Nenhum comentário neste caso.")
        
        # Adicionar rodapé personalizado, se fornecido
        if custom_footer:
            doc.add_paragraph()
            doc.add_paragraph(custom_footer)
        
        # Salvar o arquivo
        file_name = f"{slugify(case['title'])}-{report_id}.docx"
        file_path = cls.REPORTS_DIR / file_name
        
        doc.save(file_path)
        
        # Retornar o caminho relativo e o tamanho do arquivo
        file_size = os.path.getsize(file_path)
        relative_path = os.path.join('reports', file_name)
        
        return relative_path, file_size
    
    @classmethod
    def _generate_pdf_report(
        cls,
        case_data: Dict,
        template: Optional[ReportTemplate],
        custom_header: Optional[str],
        custom_footer: Optional[str],
        report_id: uuid.UUID
    ) -> tuple:
        """
        Gera um relatório em formato PDF
        """
        case = case_data['case']
        
        # Primeiro gerar o conteúdo em Markdown
        if template and template.template_content:
            # Usar o template fornecido
            md_content = template.template_content
        else:
            # Usar um template padrão
            md_content = cls._get_default_markdown_template(case_data)
        
        # Adicionar cabeçalho e rodapé personalizados, se fornecidos
        if custom_header:
            md_content = f"{custom_header}\n\n{md_content}"
        if custom_footer:
            md_content = f"{md_content}\n\n{custom_footer}"
            
        # Renderizar o template com Jinja2
        j2_env = jinja2.Environment()
        j2_template = j2_env.from_string(md_content)
        rendered_md = j2_template.render(**case_data)
        
        # Converter Markdown para HTML
        html_content = markdown.markdown(rendered_md, extensions=['tables', 'fenced_code'])
        
        # Adicionar estilos CSS para melhorar a aparência
        html_with_style = f"""
        <!DOCTYPE html>
        <html>
        <head>
            <meta charset="UTF-8">
            <title>Relatório de Caso: {case['title']}</title>
            <style>
                body {{ font-family: Arial, sans-serif; margin: 20px; }}
                h1 {{ color: #2c3e50; }}
                h2 {{ color: #3498db; border-bottom: 1px solid #eee; padding-bottom: 5px; }}
                table {{ border-collapse: collapse; width: 100%; margin-bottom: 20px; }}
                th, td {{ border: 1px solid #ddd; padding: 8px; text-align: left; }}
                th {{ background-color: #f2f2f2; }}
                tr:nth-child(even) {{ background-color: #f9f9f9; }}
                .info {{ background-color: #d9edf7; padding: 10px; border-radius: 5px; }}
                .severity-high {{ color: #d9534f; }}
                .severity-medium {{ color: #f0ad4e; }}
                .severity-low {{ color: #5bc0de; }}
                footer {{ margin-top: 30px; border-top: 1px solid #eee; padding-top: 10px; font-size: 0.8em; color: #777; }}
            </style>
        </head>
        <body>
            {html_content}
            <footer>
                <p>Relatório gerado em {timezone.now().strftime('%d/%m/%Y %H:%M:%S')}</p>
            </footer>
        </body>
        </html>
        """
        
        # Gerar PDF com WeasyPrint
        file_name = f"{slugify(case['title'])}-{report_id}.pdf"
        file_path = cls.REPORTS_DIR / file_name
        
        html = weasyprint.HTML(string=html_with_style)
        html.write_pdf(file_path)
        
        # Retornar o caminho relativo e o tamanho do arquivo
        file_size = os.path.getsize(file_path)
        relative_path = os.path.join('reports', file_name)
        
        return relative_path, file_size
    
    @staticmethod
    def _get_default_markdown_template(case_data: Dict) -> str:
        """
        Retorna um template Markdown padrão
        """
        return """
# Relatório de Caso: {{ case.title }}

**ID do Caso:** {{ case.id }}  
**Severidade:** {{ case.severity }}  
**Status:** {{ case.status }}  
**Data de Criação:** {{ case.created_at.strftime('%d/%m/%Y %H:%M') }}  
{% if case.closed_at %}**Data de Encerramento:** {{ case.closed_at.strftime('%d/%m/%Y %H:%M') }}{% endif %}

## Descrição
{{ case.description or "Sem descrição disponível." }}

{% if observables %}
## Observáveis

| Valor | Tipo | IoC | Descrição |
|-------|------|-----|-----------|
{% for obs in observables %}| {{ obs.value }} | {{ obs.type or 'N/A' }} | {{ 'Sim' if obs.is_ioc else 'Não' }} | {{ obs.description or '' }} |
{% endfor %}
{% endif %}

{% if timeline %}
## Timeline

| Data/Hora | Tipo | Descrição |
|-----------|------|-----------|
{% for event in timeline %}| {{ event.occurred_at.strftime('%d/%m/%Y %H:%M') }} | {{ event.event_type }} | {{ event.description }} |
{% endfor %}
{% endif %}

{% if mitre_techniques %}
## Técnicas MITRE ATT&CK

| ID | Nome | Táticas |
|----|------|---------|
{% for tech in mitre_techniques %}| {{ tech.id }} | {{ tech.name }} | {{ tech.tactics|map(attribute='name')|join(', ') }} |
{% endfor %}
{% endif %}

{% if tasks %}
## Tarefas

| Título | Status | Responsável | Data de Vencimento |
|--------|--------|------------|-------------------|
{% for task in tasks %}| {{ task.title }} | {{ task.status or 'N/A' }} | {{ task.assignee or 'Não atribuído' }} | {{ task.due_date.strftime('%d/%m/%Y') if task.due_date else 'N/A' }} |
{% endfor %}
{% endif %}

{% if comments %}
## Comentários

{% for comment in comments %}
**{{ comment.user_full_name or comment.user }} - {{ comment.created_at.strftime('%d/%m/%Y %H:%M') }}:**

{{ comment.text }}

{% endfor %}
{% endif %}

{% if alerts %}
## Alertas Relacionados

| Título | Severidade | Status | Fonte | Data |
|--------|------------|--------|-------|------|
{% for alert in alerts %}| {{ alert.title }} | {{ alert.severity or 'N/A' }} | {{ alert.status or 'N/A' }} | {{ alert.source_system }} | {{ alert.created_at.strftime('%d/%m/%Y') }} |
{% endfor %}
{% endif %}

---
*Relatório gerado em {{ metadata.generated_at.strftime('%d/%m/%Y %H:%M:%S') }}*
""" 